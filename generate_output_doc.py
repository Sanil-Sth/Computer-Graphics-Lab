import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
OUTPUT_DIR = "CG-Lab-Outputs"
IMAGE_ORDER = ["output_console.png", "output_gui.png", "output_zoomed.png"]

# ─────────────────────────────────────────────
# HELPER: Add border to all 4 sides
# ─────────────────────────────────────────────
def add_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    pPr.append(pBdr)

# ─────────────────────────────────────────────
# HELPER: Set spacing
# ─────────────────────────────────────────────
def set_spacing(paragraph, before=40, after=40):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)

# ─────────────────────────────────────────────
# HELPER: Set Times New Roman font on a run
# ─────────────────────────────────────────────
def set_font(run, bold=False, size_pt=14):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Also set via XML for full compatibility
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'),    'Times New Roman')
    rPr.insert(0, rFonts)

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def generate_doc_for_lab(lab_folder):
    lab_name = lab_folder.rstrip('/')

    # Collect images in order
    images_found = []
    for img_name in IMAGE_ORDER:
        img_path = os.path.join(lab_folder, img_name)
        if os.path.exists(img_path):
            images_found.append((img_name, img_path))

    # Any extra output_*.png not in standard list
    all_outputs = glob.glob(os.path.join(lab_folder, 'output_*.png'))
    known = [n for n, _ in images_found]
    for img in sorted(all_outputs):
        img_name = os.path.basename(img)
        if img_name not in known:
            images_found.append((img_name, img))

    if not images_found:
        print(f"  [SKIP] No output images in {lab_folder}")
        return

    print(f"  [BUILD] {lab_name} -> {len(images_found)} image(s)")

    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── "OUTPUTS :" heading paragraph ──
    heading = doc.add_paragraph()
    run = heading.add_run("OUTPUTS :")
    set_font(run, bold=True, size_pt=14)
    add_border(heading)
    set_spacing(heading, before=40, after=40)

    # ── Empty spacer ──
    spacer = doc.add_paragraph()
    add_border(spacer)
    set_spacing(spacer, before=40, after=40)

    # ── Each image ──
    for i, (img_name, img_path) in enumerate(images_found):

        # "Zoomed In Figure:" label before the last/zoomed image
        if img_name == "output_zoomed.png":
            label_para = doc.add_paragraph()
            label_run = label_para.add_run("Zoomed In Figure:")
            set_font(label_run, bold=False, size_pt=12)
            add_border(label_para)
            set_spacing(label_para, before=40, after=40)

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_border(img_para)
        set_spacing(img_para, before=40, after=40)

        run = img_para.add_run()
        try:
            run.add_picture(img_path, width=Inches(5.5))
        except Exception as e:
            print(f"    [WARN] {img_path}: {e}")
            img_para.add_run(f"[Image not found: {img_name}]")

    # ── Closing spacer ──
    closing = doc.add_paragraph()
    add_border(closing)
    set_spacing(closing, before=40, after=40)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, f"{lab_name}-output_print.docx")
    doc.save(out_path)
    print(f"  [SAVED] {out_path}")

# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("Scanning for lab folders...")
    lab_folders = sorted(glob.glob("Lab-*/"))
    if not lab_folders:
        print("No lab folders found. Run from repo root.")
        exit(1)
    for lab in lab_folders:
        generate_doc_for_lab(lab)
    print("\nDone! Check CG-Lab-Outputs/")
